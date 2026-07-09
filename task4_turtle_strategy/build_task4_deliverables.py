from __future__ import annotations

import html
import json
import math
import sys
from pathlib import Path

import numpy as np
import pandas as pd
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parent))
from backtester import TurtleBacktestConfig, calculate_performance_metrics, run_turtle_backtest
from turtle_strategy import TurtleStrategyConfig


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TASK_DIR = PROJECT_ROOT / "task4_turtle_strategy"
OUTPUT_DIR = TASK_DIR / "outputs"
FIG_DIR = OUTPUT_DIR / "figures"
TASK4_DOC_DIR = PROJECT_ROOT / "TASK4"
TASK4_DOC_FIG_DIR = TASK4_DOC_DIR / "figures"
DOCX_PATH = TASK4_DOC_DIR / "程冰晖 TASK4.docx"
HTML_PATH = TASK_DIR / "turtle_strategy_report.html"
NOTEBOOK_PATH = TASK_DIR / "turtle_strategy_backtest.ipynb"

STRATEGY_CONFIG = TurtleStrategyConfig(entry_window=20, exit_window=10, n_window=20, stop_n=2.0)
BACKTEST_CONFIG = TurtleBacktestConfig(
    initial_cash=100_000.0,
    risk_per_unit=0.01,
    fee_rate=0.001,
    slippage_rate=0.0005,
    allow_fractional_shares=True,
)

DATA_DIR = PROJECT_ROOT / "data/processed/equities/daily"
DEFAULT_DATA_PATH = DATA_DIR / "smic_h_00981_HK_daily_20250703_20260703.csv"


def load_font(size: int, bold: bool = False):
    candidates = [
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial.ttf",
    ]
    for path in candidates:
        try:
            return ImageFont.truetype(path, size)
        except OSError:
            continue
    return ImageFont.load_default()


FONT_TITLE = load_font(34, True)
FONT_HEAD = load_font(24, True)
FONT_REG = load_font(19)
FONT_SMALL = load_font(16)


def discover_data_files() -> list[Path]:
    files = sorted(DATA_DIR.glob("*_daily_*.csv"))
    return [p for p in files if "three_companies" not in p.name]


def read_stock_data(path: Path) -> pd.DataFrame:
    data = pd.read_csv(path, dtype={"trade_date": str, "ts_code": str}, encoding="utf-8-sig")
    data = data.sort_values("trade_date").reset_index(drop=True)
    data["date"] = pd.to_datetime(data["trade_date"], format="%Y%m%d")
    return data


def instrument_label(data: pd.DataFrame, path: Path) -> str:
    name_map = {
        "smic_h": "中芯国际港股",
        "smic_a": "中芯国际A股",
        "byd_a": "比亚迪A股",
        "byd_h": "比亚迪港股",
        "cypc_a": "长江电力A股",
    }
    key = str(data.get("instrument_key", pd.Series([path.stem.split("_daily_")[0]])).iloc[0])
    ts_code = str(data["ts_code"].iloc[0])
    return f"{name_map.get(key, key)}（{ts_code}）"


def short_label(label: str) -> str:
    return (
        label.replace("（", "\n")
        .replace("）", "")
        .replace("中芯国际", "中芯")
        .replace("长江电力", "长电")
    )


def pct(value: float, digits: int = 2) -> str:
    if pd.isna(value):
        return "-"
    return f"{value * 100:.{digits}f}%"


def num(value: float, digits: int = 2) -> str:
    if pd.isna(value):
        return "-"
    return f"{value:,.{digits}f}"


def safe_metric(value: object) -> float | None:
    try:
        out = float(value)
    except (TypeError, ValueError):
        return None
    if math.isnan(out) or math.isinf(out):
        return None
    return out


def build_backtests() -> tuple[dict[str, pd.DataFrame], list[dict[str, object]]]:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    portfolios: dict[str, pd.DataFrame] = {}
    metrics_rows: list[dict[str, object]] = []

    for path in discover_data_files():
        data = read_stock_data(path)
        label = instrument_label(data, path)
        portfolio = run_turtle_backtest(data, STRATEGY_CONFIG, BACKTEST_CONFIG)
        metrics = calculate_performance_metrics(portfolio, BACKTEST_CONFIG)
        key = str(data["instrument_key"].iloc[0])
        filename_ts = str(data["ts_code"].iloc[0]).replace(".", "_")
        out_path = OUTPUT_DIR / f"{key}_{filename_ts}_turtle_core_backtest.csv"
        portfolio.to_csv(out_path, index=False, encoding="utf-8-sig")
        metrics.update(
            {
                "instrument_key": key,
                "label": label,
                "ts_code": str(data["ts_code"].iloc[0]),
                "market": str(data["market"].iloc[0]),
                "currency": str(data["currency"].iloc[0]),
                "rows": int(len(data)),
                "start_date": data["date"].dt.strftime("%Y-%m-%d").iloc[0],
                "end_date": data["date"].dt.strftime("%Y-%m-%d").iloc[-1],
                "data_file": str(path.relative_to(PROJECT_ROOT)),
                "backtest_file": str(out_path.relative_to(PROJECT_ROOT)),
            }
        )
        portfolios[key] = portfolio
        metrics_rows.append(metrics)

    summary = pd.DataFrame(metrics_rows)
    summary_path = OUTPUT_DIR / "turtle_core_metrics_summary.csv"
    summary.to_csv(summary_path, index=False, encoding="utf-8-sig")
    json_path = OUTPUT_DIR / "turtle_core_metrics_summary.json"
    json_path.write_text(json.dumps(metrics_rows, ensure_ascii=False, indent=2), encoding="utf-8")
    return portfolios, metrics_rows


def parameter_sensitivity(default_path: Path = DEFAULT_DATA_PATH) -> pd.DataFrame:
    data = read_stock_data(default_path)
    rows: list[dict[str, object]] = []
    for entry, exit_ in [(20, 10), (30, 15), (40, 20), (55, 20)]:
        cfg = TurtleStrategyConfig(entry_window=entry, exit_window=exit_, n_window=20, stop_n=2.0)
        portfolio = run_turtle_backtest(data, cfg, BACKTEST_CONFIG)
        metrics = calculate_performance_metrics(portfolio, BACKTEST_CONFIG)
        rows.append(
            {
                "entry_window": entry,
                "exit_window": exit_,
                "parameter": f"{entry}/{exit_}",
                "cumulative_return": metrics["cumulative_return"],
                "max_drawdown": metrics["max_drawdown"],
                "sharpe_ratio": metrics["sharpe_ratio"],
                "trade_count": metrics["trade_count"],
            }
        )
    sens = pd.DataFrame(rows)
    sens.to_csv(OUTPUT_DIR / "turtle_parameter_sensitivity.csv", index=False, encoding="utf-8-sig")
    return sens


def chart_canvas(title: str, subtitle: str = "", width: int = 1400, height: int = 850):
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    draw.text((width / 2, 28), title, fill="#172033", font=FONT_TITLE, anchor="ma")
    if subtitle:
        draw.text((width / 2, 72), subtitle, fill="#556070", font=FONT_REG, anchor="ma")
    return img, draw


def scale(values, lo, hi, low_px, high_px):
    vals = np.asarray(values, dtype=float)
    if not np.isfinite(lo) or not np.isfinite(hi) or hi == lo:
        return np.full(len(vals), (low_px + high_px) / 2)
    return low_px - (vals - lo) / (hi - lo) * (low_px - high_px)


def draw_axes(draw, rect, y_min, y_max, pct_axis=False):
    left, top, right, bottom = rect
    draw.rectangle(rect, outline="#D5DCE6", width=1)
    for i in range(5):
        y = int(top + (bottom - top) * i / 4)
        value = y_max - (y_max - y_min) * i / 4
        label = pct(value, 0) if pct_axis else f"{value:.1f}"
        draw.line((left, y, right, y), fill="#ECF0F5", width=1)
        draw.text((left - 10, y), label, fill="#5B6675", font=FONT_SMALL, anchor="rm")


def draw_line(draw, x_values, y_values, color, width=3):
    pts = [(int(x), int(y)) for x, y in zip(x_values, y_values) if np.isfinite(y)]
    if len(pts) > 1:
        draw.line(pts, fill=color, width=width, joint="curve")


def draw_x_labels(draw, dates, left, right, y):
    if len(dates) == 0:
        return
    positions = [0, len(dates) // 2, len(dates) - 1]
    for idx in positions:
        x = left + (right - left) * idx / max(1, len(dates) - 1)
        label = pd.to_datetime(dates.iloc[idx]).strftime("%Y-%m-%d")
        draw.text((x, y), label, fill="#5B6675", font=FONT_SMALL, anchor="ma")


def figure_turtle_signals(portfolio: pd.DataFrame, label: str, path: Path) -> None:
    img, draw = chart_canvas(
        f"图1 {label} 海龟策略通道与交易信号",
        "复权收盘价、20日高点通道、10日低点通道，信号为收盘确认后下一交易日执行",
    )
    rect = (105, 125, 1315, 720)
    df = portfolio.tail(170).copy()
    price_cols = ["qfq_close", "entry_high_20", "exit_low_10"]
    values = pd.concat([df[col] for col in price_cols], ignore_index=True).dropna()
    y_min, y_max = float(values.min()), float(values.max())
    pad = (y_max - y_min) * 0.08
    y_min -= pad
    y_max += pad
    draw_axes(draw, rect, y_min, y_max)
    left, top, right, bottom = rect
    x = np.linspace(left, right, len(df))
    for col, color, name in [
        ("qfq_close", "#213547", "qfq_close"),
        ("entry_high_20", "#D97706", "20日入场通道"),
        ("exit_low_10", "#0F766E", "10日退出通道"),
    ]:
        y = scale(df[col], y_min, y_max, bottom, top)
        draw_line(draw, x, y, color, 3 if col == "qfq_close" else 2)
        draw.text((right - 180, top + 28 + 28 * price_cols.index(col)), name, fill=color, font=FONT_SMALL)

    y_price = scale(df["qfq_close"], y_min, y_max, bottom, top)
    for idx, row in df.reset_index(drop=True).iterrows():
        action = row["trade_action"]
        if action == "BUY":
            xx, yy = int(x[idx]), int(y_price[idx])
            draw.polygon([(xx, yy - 14), (xx - 10, yy + 8), (xx + 10, yy + 8)], fill="#16A34A")
        elif action == "SELL":
            xx, yy = int(x[idx]), int(y_price[idx])
            draw.polygon([(xx, yy + 14), (xx - 10, yy - 8), (xx + 10, yy - 8)], fill="#DC2626")

    draw_x_labels(draw, df["date"], left, right, bottom + 34)
    draw.text((left, 790), "绿色三角表示买入，红色三角表示卖出；通道整体 shift(1)，不使用当日未来信息。", fill="#4B5563", font=FONT_REG)
    img.save(path)


def figure_equity_drawdown(portfolio: pd.DataFrame, label: str, path: Path) -> None:
    img, draw = chart_canvas(f"图2 {label} 策略净值与回撤", "Turtle Core 20/10 vs 买入持有基准")
    df = portfolio.copy()
    top_rect = (105, 125, 1315, 475)
    bottom_rect = (105, 560, 1315, 760)
    equity_vals = pd.concat([df["portfolio_value"], df["benchmark_value"]])
    y_min, y_max = float(equity_vals.min()), float(equity_vals.max())
    pad = (y_max - y_min) * 0.08
    draw_axes(draw, top_rect, y_min - pad, y_max + pad)
    x = np.linspace(top_rect[0], top_rect[2], len(df))
    draw_line(draw, x, scale(df["portfolio_value"], y_min - pad, y_max + pad, top_rect[3], top_rect[1]), "#1D4ED8", 3)
    draw_line(draw, x, scale(df["benchmark_value"], y_min - pad, y_max + pad, top_rect[3], top_rect[1]), "#64748B", 3)
    draw.text((1120, 150), "策略净值", fill="#1D4ED8", font=FONT_SMALL)
    draw.text((1120, 180), "买入持有", fill="#64748B", font=FONT_SMALL)

    dd_min = min(float(df["drawdown"].min()), -0.01)
    draw_axes(draw, bottom_rect, dd_min, 0.0, pct_axis=True)
    y_dd = scale(df["drawdown"], dd_min, 0.0, bottom_rect[3], bottom_rect[1])
    pts = [(int(xx), int(yy)) for xx, yy in zip(x, y_dd)]
    baseline = bottom_rect[1]
    if len(pts) > 1:
        poly = [(pts[0][0], baseline), *pts, (pts[-1][0], baseline)]
        draw.polygon(poly, fill="#F5CBD2")
        draw.line(pts, fill="#BE123C", width=2)
    draw_x_labels(draw, df["date"], top_rect[0], top_rect[2], bottom_rect[3] + 34)
    img.save(path)


def figure_multi_metrics(metrics_rows: list[dict[str, object]], path: Path) -> None:
    img, draw = chart_canvas("图3 多标的海龟策略绩效对比", "同一 Turtle Core 20/10 规则在五只股票上的历史表现")
    rows = metrics_rows
    labels = [short_label(str(r["label"])) for r in rows]
    metrics = [
        ("cumulative_return", "累计收益", "#1D4ED8"),
        ("max_drawdown", "最大回撤", "#BE123C"),
        ("sharpe_ratio", "Sharpe", "#0F766E"),
    ]
    panel_w = 380
    base_x = 130
    for p_idx, (key, title, color) in enumerate(metrics):
        left = base_x + p_idx * (panel_w + 28)
        top, bottom = 135, 665
        draw.text((left + panel_w / 2, 112), title, fill="#172033", font=FONT_HEAD, anchor="ma")
        vals = [safe_metric(r.get(key)) or 0.0 for r in rows]
        if key == "max_drawdown":
            vals = [abs(v) for v in vals]
        scaled_vals = [abs(v) for v in vals]
        max_v = max(max(scaled_vals), 0.01)
        bar_w = 44
        gap = (panel_w - len(vals) * bar_w) / (len(vals) + 1)
        for idx, val in enumerate(vals):
            x0 = left + gap + idx * (bar_w + gap)
            h = (abs(val) / max_v) * (bottom - top)
            y0 = bottom - h
            draw.rectangle((x0, y0, x0 + bar_w, bottom), fill=color)
            label = f"{val:.2f}" if key == "sharpe_ratio" else pct(val, 0)
            draw.text((x0 + bar_w / 2, y0 - 18), label, fill="#334155", font=FONT_SMALL, anchor="ma")
            draw.text((x0 + bar_w / 2, bottom + 12), labels[idx], fill="#334155", font=FONT_SMALL, anchor="ma")
        draw.line((left, bottom, left + panel_w, bottom), fill="#CBD5E1", width=1)
    draw.text((130, 760), "说明：最大回撤图以绝对值展示，越低越好；Sharpe 只用于样本内横向比较。", fill="#4B5563", font=FONT_REG)
    img.save(path)


def figure_parameter_sensitivity(sens: pd.DataFrame, path: Path) -> None:
    img, draw = chart_canvas("图4 核心参数敏感性分析", "默认以中芯国际港股为例，比较不同入场/退出通道组合")
    left, top, right, bottom = 170, 145, 1230, 625
    vals = sens["cumulative_return"].astype(float).to_numpy()
    dd = sens["max_drawdown"].astype(float).to_numpy()
    max_abs = max(abs(vals).max(), 0.01)
    cell_w = (right - left) / len(sens)
    for idx, row in sens.iterrows():
        val = float(row["cumulative_return"])
        intensity = min(1.0, abs(val) / max_abs)
        if val >= 0:
            base = np.array([39, 174, 96])
        else:
            base = np.array([192, 57, 43])
        color = tuple((255 - (255 - base) * (0.25 + 0.65 * intensity)).astype(int))
        x0, x1 = left + idx * cell_w, left + (idx + 1) * cell_w - 10
        draw.rounded_rectangle((x0, top, x1, bottom), radius=14, fill=color, outline="#FFFFFF", width=3)
        draw.text(((x0 + x1) / 2, top + 52), str(row["parameter"]), fill="#172033", font=FONT_HEAD, anchor="ma")
        draw.text(((x0 + x1) / 2, top + 150), f"累计收益 {pct(val)}", fill="#172033", font=FONT_REG, anchor="ma")
        draw.text(((x0 + x1) / 2, top + 210), f"最大回撤 {pct(float(row['max_drawdown']))}", fill="#172033", font=FONT_REG, anchor="ma")
        sharpe = safe_metric(row["sharpe_ratio"])
        draw.text(((x0 + x1) / 2, top + 270), f"Sharpe {num(sharpe) if sharpe is not None else '-'}", fill="#172033", font=FONT_REG, anchor="ma")
        draw.text(((x0 + x1) / 2, top + 330), f"交易次数 {int(row['trade_count'])}", fill="#172033", font=FONT_REG, anchor="ma")
    draw.text((left, 710), "颜色深度表示累计收益绝对值；参数越长，信号通常越少，对趋势持续性要求更高。", fill="#4B5563", font=FONT_REG)
    img.save(path)


def copy_figures_to_task4() -> None:
    TASK4_DOC_FIG_DIR.mkdir(parents=True, exist_ok=True)
    for src in FIG_DIR.glob("fig*.png"):
        (TASK4_DOC_FIG_DIR / src.name).write_bytes(src.read_bytes())


def build_figures(portfolios: dict[str, pd.DataFrame], metrics_rows: list[dict[str, object]], sens: pd.DataFrame) -> None:
    default_key = "smic_h" if "smic_h" in portfolios else next(iter(portfolios))
    default_label = next(r["label"] for r in metrics_rows if r["instrument_key"] == default_key)
    figure_turtle_signals(portfolios[default_key], str(default_label), FIG_DIR / "fig1_turtle_signals.png")
    figure_equity_drawdown(portfolios[default_key], str(default_label), FIG_DIR / "fig2_equity_drawdown.png")
    figure_multi_metrics(metrics_rows, FIG_DIR / "fig3_multi_stock_metrics.png")
    figure_parameter_sensitivity(sens, FIG_DIR / "fig4_parameter_sensitivity.png")
    copy_figures_to_task4()


def build_html_report(portfolios: dict[str, pd.DataFrame], metrics_rows: list[dict[str, object]], sens: pd.DataFrame) -> None:
    datasets = []
    for row in metrics_rows:
        key = str(row["instrument_key"])
        p = portfolios[key]
        sampled = p[["date", "qfq_close", "entry_high_20", "exit_low_10", "portfolio_value", "benchmark_value", "drawdown", "trade_action"]].copy()
        datasets.append(
            {
                "key": key,
                "label": row["label"],
                "metrics": {k: safe_metric(v) if isinstance(v, (int, float, np.floating)) else v for k, v in row.items()},
                "rows": [
                    {
                        "date": r.date.strftime("%Y-%m-%d"),
                        "close": safe_metric(r.qfq_close),
                        "entryHigh": safe_metric(r.entry_high_20),
                        "exitLow": safe_metric(r.exit_low_10),
                        "portfolio": safe_metric(r.portfolio_value),
                        "benchmark": safe_metric(r.benchmark_value),
                        "drawdown": safe_metric(r.drawdown),
                        "tradeAction": str(r.trade_action),
                    }
                    for r in sampled.itertuples(index=False)
                ],
            }
        )

    payload = json.dumps({"datasets": datasets, "sensitivity": sens.to_dict("records")}, ensure_ascii=False)
    metric_rows = "\n".join(
        "<tr>"
        f"<td>{html.escape(str(r['label']))}</td>"
        f"<td>{pct(float(r['cumulative_return']))}</td>"
        f"<td>{pct(float(r['max_drawdown']))}</td>"
        f"<td>{num(float(r['sharpe_ratio']))}</td>"
        f"<td>{pct(float(r['benchmark_cumulative_return']))}</td>"
        f"<td>{int(r['trade_count'])}</td>"
        "</tr>"
        for r in metrics_rows
    )
    HTML_PATH.write_text(
        f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Task 4 海龟策略回测</title>
  <style>
    :root {{ --ink:#172033; --muted:#64748b; --line:#d8dee8; --bg:#f6f8fb; --panel:#ffffff; --blue:#1d4ed8; --green:#15803d; --red:#be123c; --amber:#d97706; }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; font-family:-apple-system,BlinkMacSystemFont,"Segoe UI","PingFang SC","Microsoft YaHei",sans-serif; background:var(--bg); color:var(--ink); }}
    header {{ padding:28px 36px 18px; background:#fff; border-bottom:1px solid var(--line); display:flex; justify-content:space-between; gap:24px; align-items:flex-end; }}
    h1 {{ margin:0 0 8px; font-size:34px; letter-spacing:0; }}
    h2 {{ margin:0 0 12px; font-size:20px; }}
    p {{ margin:0; color:var(--muted); line-height:1.65; }}
    main {{ width:min(1280px, calc(100vw - 32px)); margin:22px auto 42px; display:grid; gap:18px; }}
    .toolbar, .panel {{ background:#fff; border:1px solid var(--line); border-radius:8px; padding:18px; }}
    .toolbar {{ display:flex; flex-wrap:wrap; gap:14px; align-items:center; justify-content:space-between; }}
    select {{ font:inherit; padding:9px 12px; border:1px solid var(--line); border-radius:6px; background:white; min-width:230px; }}
    .kpis {{ display:grid; grid-template-columns:repeat(5, minmax(130px,1fr)); gap:12px; }}
    .kpi {{ background:#fff; border:1px solid var(--line); border-radius:8px; padding:14px; }}
    .kpi span {{ display:block; color:var(--muted); font-size:13px; margin-bottom:6px; }}
    .kpi strong {{ font-size:22px; }}
    .grid {{ display:grid; grid-template-columns:1.35fr .85fr; gap:18px; }}
    svg {{ width:100%; height:390px; display:block; }}
    table {{ width:100%; border-collapse:collapse; font-size:14px; }}
    th, td {{ border-bottom:1px solid var(--line); padding:10px 8px; text-align:right; }}
    th:first-child, td:first-child {{ text-align:left; }}
    th {{ color:#334155; background:#f8fafc; }}
    .small {{ font-size:13px; color:var(--muted); }}
    .pill {{ display:inline-flex; align-items:center; border:1px solid var(--line); border-radius:999px; padding:6px 10px; color:#334155; background:#fff; font-size:13px; }}
    .figure-links {{ display:grid; grid-template-columns:repeat(4,1fr); gap:12px; }}
    .figure-links a {{ display:block; border:1px solid var(--line); border-radius:8px; padding:12px; color:var(--blue); text-decoration:none; background:#fff; }}
    @media (max-width:900px) {{ header,.grid {{ grid-template-columns:1fr; display:block; }} .kpis,.figure-links {{ grid-template-columns:1fr 1fr; }} header > * + * {{ margin-top:14px; }} }}
  </style>
</head>
<body>
  <header>
    <div>
      <p class="small">Task 4 · Turtle Core 20/10 · qfq adjusted prices</p>
      <h1>海龟策略回测工作台</h1>
      <p>展示唐奇安通道、ATR/N 值、买卖信号、风险定仓和样本内回测表现。内容仅用于课程研究，不构成投资建议。</p>
    </div>
    <a class="pill" href="../index.html">返回任务汇总页</a>
  </header>
  <main>
    <section class="toolbar">
      <label>选择股票数据 <select id="dataset"></select></label>
      <span class="pill">入场通道 20 日</span>
      <span class="pill">退出通道 10 日</span>
      <span class="pill">N 值 20 日</span>
      <span class="pill">2N 止损</span>
    </section>
    <section class="kpis">
      <div class="kpi"><span>累计收益</span><strong id="cum">-</strong></div>
      <div class="kpi"><span>最大回撤</span><strong id="dd">-</strong></div>
      <div class="kpi"><span>Sharpe</span><strong id="sharpe">-</strong></div>
      <div class="kpi"><span>买入持有</span><strong id="bench">-</strong></div>
      <div class="kpi"><span>交易次数</span><strong id="trades">-</strong></div>
    </section>
    <section class="grid">
      <div class="panel">
        <h2 id="chartTitle">价格、通道与交易信号</h2>
        <svg id="priceSvg" role="img" aria-label="价格通道图"></svg>
      </div>
      <div class="panel">
        <h2>净值与回撤</h2>
        <svg id="equitySvg" role="img" aria-label="净值回撤图"></svg>
      </div>
    </section>
    <section class="panel">
      <h2>多标的绩效汇总</h2>
      <table>
        <thead><tr><th>标的</th><th>累计收益</th><th>最大回撤</th><th>Sharpe</th><th>买入持有</th><th>交易次数</th></tr></thead>
        <tbody>{metric_rows}</tbody>
      </table>
    </section>
    <section class="panel">
      <h2>参数敏感性</h2>
      <svg id="sensSvg" role="img" aria-label="参数敏感性图"></svg>
      <p class="small">参数组合格式为“入场窗口/退出窗口”。更长通道通常减少信号数量，更适合持续趋势；短通道更敏感，也更容易受到震荡噪声影响。</p>
    </section>
    <section class="panel">
      <h2>报告图表文件</h2>
      <div class="figure-links">
        <a href="./outputs/figures/fig1_turtle_signals.png">图1 通道与信号</a>
        <a href="./outputs/figures/fig2_equity_drawdown.png">图2 净值与回撤</a>
        <a href="./outputs/figures/fig3_multi_stock_metrics.png">图3 多标的绩效</a>
        <a href="./outputs/figures/fig4_parameter_sensitivity.png">图4 参数敏感性</a>
      </div>
    </section>
  </main>
  <script>
    const payload = {payload};
    const pct = v => v === null || Number.isNaN(v) ? "-" : (v * 100).toFixed(2) + "%";
    const num = v => v === null || Number.isNaN(v) ? "-" : Number(v).toFixed(2);
    const dataset = document.querySelector("#dataset");
    payload.datasets.forEach((d, i) => {{
      const opt = document.createElement("option");
      opt.value = i;
      opt.textContent = d.label;
      dataset.appendChild(opt);
    }});
    function points(rows, key, w, h, pad, minV, maxV) {{
      return rows.map((r, i) => {{
        const x = pad + i * (w - 2 * pad) / Math.max(1, rows.length - 1);
        const v = r[key];
        const y = h - pad - (v - minV) / Math.max(1e-9, maxV - minV) * (h - 2 * pad);
        return [x, y, v];
      }}).filter(p => p[2] !== null && !Number.isNaN(p[2]));
    }}
    function pathFrom(pts) {{ return pts.map((p,i)=>(i?'L':'M')+p[0].toFixed(1)+','+p[1].toFixed(1)).join(' '); }}
    function drawPrice(d) {{
      const svg = document.querySelector("#priceSvg");
      const rows = d.rows;
      const w = svg.clientWidth || 760, h = 390, pad = 42;
      const values = rows.flatMap(r => [r.close, r.entryHigh, r.exitLow]).filter(v => v !== null);
      const minV = Math.min(...values), maxV = Math.max(...values), span = maxV - minV || 1;
      const lo = minV - span * .08, hi = maxV + span * .08;
      const close = points(rows, 'close', w, h, pad, lo, hi);
      const entry = points(rows, 'entryHigh', w, h, pad, lo, hi);
      const exit = points(rows, 'exitLow', w, h, pad, lo, hi);
      const marks = rows.map((r,i)=>({{...r, x:pad+i*(w-2*pad)/Math.max(1,rows.length-1)}})).filter(r=>r.tradeAction==='BUY'||r.tradeAction==='SELL');
      svg.setAttribute('viewBox', `0 0 ${{w}} ${{h}}`);
      svg.innerHTML = `<rect x="${{pad}}" y="20" width="${{w-2*pad}}" height="${{h-2*pad}}" fill="white" stroke="#d8dee8"/>
        <path d="${{pathFrom(exit)}}" fill="none" stroke="#0f766e" stroke-width="2"/>
        <path d="${{pathFrom(entry)}}" fill="none" stroke="#d97706" stroke-width="2"/>
        <path d="${{pathFrom(close)}}" fill="none" stroke="#213547" stroke-width="2.5"/>
        ${{marks.map(m=>`<circle cx="${{m.x}}" cy="${{points([m],'close',w,h,pad,lo,hi)[0]?.[1] ?? 0}}" r="5" fill="${{m.tradeAction==='BUY'?'#16a34a':'#dc2626'}}"/>`).join('')}}
        <text x="${{pad}}" y="${{h-8}}" fill="#64748b" font-size="12">${{rows[0].date}}</text>
        <text x="${{w-pad}}" y="${{h-8}}" fill="#64748b" font-size="12" text-anchor="end">${{rows[rows.length-1].date}}</text>`;
    }}
    function drawEquity(d) {{
      const svg = document.querySelector("#equitySvg");
      const rows = d.rows;
      const w = svg.clientWidth || 520, h = 390, pad = 42;
      const values = rows.flatMap(r => [r.portfolio, r.benchmark]).filter(v => v !== null);
      const minV = Math.min(...values), maxV = Math.max(...values), span = maxV - minV || 1;
      const port = points(rows, 'portfolio', w, 245, pad, minV-span*.08, maxV+span*.08);
      const bench = points(rows, 'benchmark', w, 245, pad, minV-span*.08, maxV+span*.08);
      const minDd = Math.min(...rows.map(r=>r.drawdown).filter(v=>v!==null), -0.01);
      const dd = rows.map((r,i)=>[pad+i*(w-2*pad)/Math.max(1,rows.length-1), 350 - (r.drawdown - minDd)/(0-minDd)*(92)]).filter(p=>Number.isFinite(p[1]));
      svg.setAttribute('viewBox', `0 0 ${{w}} ${{h}}`);
      svg.innerHTML = `<rect x="${{pad}}" y="20" width="${{w-2*pad}}" height="210" fill="white" stroke="#d8dee8"/>
        <path d="${{pathFrom(bench)}}" fill="none" stroke="#64748b" stroke-width="2"/>
        <path d="${{pathFrom(port)}}" fill="none" stroke="#1d4ed8" stroke-width="2.5"/>
        <rect x="${{pad}}" y="258" width="${{w-2*pad}}" height="92" fill="white" stroke="#d8dee8"/>
        <path d="${{pathFrom(dd)}}" fill="none" stroke="#be123c" stroke-width="2"/>
        <text x="${{pad}}" y="376" fill="#64748b" font-size="12">蓝线为策略净值，灰线为买入持有，红线为回撤</text>`;
    }}
    function drawSensitivity() {{
      const svg = document.querySelector("#sensSvg");
      const w = svg.clientWidth || 1160, h = 250, pad = 38;
      const rows = payload.sensitivity;
      const maxAbs = Math.max(.01, ...rows.map(r=>Math.abs(r.cumulative_return)));
      const cell = (w - 2*pad) / rows.length;
      svg.setAttribute('viewBox', `0 0 ${{w}} ${{h}}`);
      svg.innerHTML = rows.map((r,i)=>{{
        const x = pad + i*cell + 6;
        const fill = r.cumulative_return >= 0 ? '#bfe6ce' : '#f2c4c4';
        return `<rect x="${{x}}" y="30" width="${{cell-12}}" height="150" rx="8" fill="${{fill}}" stroke="#d8dee8"/>
          <text x="${{x+(cell-12)/2}}" y="66" text-anchor="middle" font-size="17" fill="#172033">${{r.parameter}}</text>
          <text x="${{x+(cell-12)/2}}" y="105" text-anchor="middle" font-size="14" fill="#172033">收益 ${{pct(r.cumulative_return)}}</text>
          <text x="${{x+(cell-12)/2}}" y="135" text-anchor="middle" font-size="14" fill="#172033">回撤 ${{pct(r.max_drawdown)}}</text>
          <text x="${{x+(cell-12)/2}}" y="165" text-anchor="middle" font-size="14" fill="#172033">交易 ${{r.trade_count}}</text>`;
      }}).join('');
    }}
    function render() {{
      const d = payload.datasets[Number(dataset.value)];
      document.querySelector("#chartTitle").textContent = d.label + "：价格、通道与交易信号";
      document.querySelector("#cum").textContent = pct(d.metrics.cumulative_return);
      document.querySelector("#dd").textContent = pct(d.metrics.max_drawdown);
      document.querySelector("#sharpe").textContent = num(d.metrics.sharpe_ratio);
      document.querySelector("#bench").textContent = pct(d.metrics.benchmark_cumulative_return);
      document.querySelector("#trades").textContent = d.metrics.trade_count;
      drawPrice(d); drawEquity(d); drawSensitivity();
    }}
    dataset.addEventListener('change', render);
    window.addEventListener('resize', render);
    render();
  </script>
</body>
</html>
""",
        encoding="utf-8",
    )
    (TASK_DIR / "index.html").write_text(
        """<!doctype html>
<html lang="zh-CN">
<head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><meta http-equiv="refresh" content="0; url=turtle_strategy_report.html"><title>Task 4 海龟策略回测</title></head>
<body><p><a href="turtle_strategy_report.html">打开 Task 4 海龟策略回测报告</a></p></body>
</html>
""",
        encoding="utf-8",
    )


def build_interactive_html_report() -> None:
    datasets = []
    for path in discover_data_files():
        data = read_stock_data(path)
        label = instrument_label(data, path)
        datasets.append(
            {
                "key": str(data["instrument_key"].iloc[0]),
                "label": label,
                "tsCode": str(data["ts_code"].iloc[0]),
                "market": str(data["market"].iloc[0]),
                "currency": str(data["currency"].iloc[0]),
                "startDate": data["date"].dt.strftime("%Y-%m-%d").iloc[0],
                "endDate": data["date"].dt.strftime("%Y-%m-%d").iloc[-1],
                "rows": [
                    {
                        "date": row.date.strftime("%Y-%m-%d"),
                        "open": safe_metric(row.qfq_open),
                        "high": safe_metric(row.qfq_high),
                        "low": safe_metric(row.qfq_low),
                        "close": safe_metric(row.qfq_close),
                        "preClose": safe_metric(row.qfq_pre_close),
                    }
                    for row in data[["date", "qfq_open", "qfq_high", "qfq_low", "qfq_close", "qfq_pre_close"]].itertuples(index=False)
                ],
            }
        )
    payload = json.dumps({"datasets": datasets}, ensure_ascii=False)
    template = r"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>Task 4 海龟策略回测</title>
  <style>
    :root { --bg:#f5f7fa; --surface:#fff; --line:#d9e0e8; --text:#1d2733; --muted:#667385; --blue:#1f6fb2; --orange:#d9822b; --green:#2e8f57; --red:#b94d4d; --shadow:0 10px 30px rgba(28,39,51,.08); }
    * { box-sizing:border-box; }
    body { margin:0; background:var(--bg); color:var(--text); font-family:Inter, ui-sans-serif, system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC", "Microsoft YaHei", sans-serif; }
    button, input { font:inherit; }
    .topbar { display:flex; align-items:center; justify-content:space-between; gap:24px; padding:18px 22px; background:var(--surface); border-bottom:1px solid var(--line); position:sticky; top:0; z-index:5; }
    .eyebrow { margin:0 0 4px; font-size:12px; color:var(--muted); letter-spacing:0; text-transform:uppercase; }
    h1, h2, h3 { margin:0; letter-spacing:0; }
    h1 { font-size:24px; line-height:1.15; }
    h2 { font-size:18px; }
    h3 { font-size:14px; }
    p { margin:0; color:var(--muted); line-height:1.6; }
    .top-actions { display:flex; align-items:end; gap:10px; flex-wrap:wrap; justify-content:flex-end; }
    .select-label { display:grid; gap:5px; color:var(--muted); font-size:12px; }
    .select-label input { min-width:150px; height:36px; border:1px solid var(--line); border-radius:6px; padding:0 10px; background:#fff; color:var(--text); }
    .command-btn { height:36px; border:1px solid #245f99; border-radius:6px; background:var(--blue); color:#fff; padding:0 12px; cursor:pointer; }
    .command-btn.secondary { color:var(--blue); background:#f7fbff; }
    .workspace { display:grid; grid-template-columns:260px minmax(520px,1fr) 330px; gap:14px; padding:14px; }
    .data-panel, .chart-panel, .param-panel { background:var(--surface); border:1px solid var(--line); border-radius:8px; box-shadow:var(--shadow); }
    .data-panel, .param-panel { padding:14px; align-self:start; position:sticky; top:84px; max-height:calc(100vh - 104px); overflow:auto; }
    .chart-panel { min-width:0; padding:16px; }
    .panel-heading, .chart-header, .chart-block-head { display:flex; align-items:center; justify-content:space-between; gap:12px; }
    .muted { color:var(--muted); font-size:12px; }
    .module-list { display:grid; gap:10px; margin-top:14px; }
    .data-card { width:100%; text-align:left; border:1px solid var(--line); border-radius:8px; background:#fbfcfe; padding:11px; cursor:pointer; }
    .data-card.active { border-color:var(--blue); box-shadow:inset 3px 0 0 var(--blue); background:#f5f9ff; }
    .data-card strong { display:block; font-size:14px; margin-bottom:4px; }
    .data-card span { display:block; color:var(--muted); font-size:12px; line-height:1.45; }
    .pill-row { display:flex; flex-wrap:wrap; justify-content:flex-end; gap:6px; }
    .pill { border:1px solid var(--line); border-radius:999px; padding:4px 8px; font-size:12px; color:var(--muted); background:#fbfcfe; }
    .warning-box { margin-top:12px; padding:10px 12px; border:1px solid #f1c27d; background:#fff8ed; color:#8a5a0a; border-radius:8px; font-size:13px; }
    .hidden { display:none; }
    .summary-strip { display:grid; grid-template-columns:repeat(5,minmax(120px,1fr)); gap:10px; margin-top:14px; }
    .summary-strip div { border:1px solid var(--line); border-radius:8px; background:#fbfcfe; padding:10px; }
    .summary-label { display:block; color:var(--muted); font-size:12px; margin-bottom:4px; }
    .summary-strip strong { font-size:18px; }
    .chart-grid { display:grid; gap:14px; margin-top:14px; }
    .chart-block { border:1px solid var(--line); border-radius:8px; padding:12px; background:#fff; min-width:0; }
    .chart-block.large svg { height:420px; }
    .chart-block svg { width:100%; height:300px; display:block; }
    .table-panel { margin-top:14px; border:1px solid var(--line); border-radius:8px; overflow:auto; background:#fff; }
    table { width:100%; border-collapse:collapse; font-size:13px; }
    th, td { border-bottom:1px solid var(--line); padding:9px 8px; text-align:right; white-space:nowrap; }
    th:first-child, td:first-child { text-align:left; }
    th { background:#f8fafc; color:#334155; }
    .control-section { border-top:1px solid var(--line); padding-top:12px; margin-top:12px; }
    .control-section:first-of-type { border-top:0; padding-top:0; }
    .control-section h3 { margin-bottom:10px; }
    .control-row { display:grid; grid-template-columns:1fr 96px; align-items:center; gap:8px; margin:10px 0; color:var(--muted); font-size:12px; }
    .control-row input[type=number] { width:96px; height:32px; border:1px solid var(--line); border-radius:6px; padding:0 8px; color:var(--text); }
    .toggle-grid { display:grid; grid-template-columns:1fr 1fr; gap:8px; color:var(--muted); font-size:13px; }
    .preset-grid { display:grid; grid-template-columns:1fr 1fr; gap:8px; }
    .preset-grid button { height:34px; border:1px solid var(--line); border-radius:6px; background:#fbfcfe; cursor:pointer; color:var(--text); }
    .preset-grid button.active { border-color:var(--blue); color:var(--blue); background:#f5f9ff; }
    .small-copy { font-size:12px; color:var(--muted); line-height:1.55; }
    @media (max-width:1100px) { .workspace { grid-template-columns:1fr; } .data-panel, .param-panel { position:static; max-height:none; } .module-list { grid-template-columns:repeat(2,minmax(0,1fr)); } }
    @media (max-width:720px) { .topbar { display:block; } .top-actions { justify-content:flex-start; margin-top:12px; } .module-list, .summary-strip { grid-template-columns:1fr; } }
  </style>
</head>
<body>
  <div class="app-shell">
    <header class="topbar">
      <div>
        <p class="eyebrow">Task 4 · Turtle Strategy Lab</p>
        <h1>海龟策略回测工作台</h1>
      </div>
      <div class="top-actions">
        <label class="select-label"><span>日期起点</span><input id="startDate" type="date"></label>
        <label class="select-label"><span>日期终点</span><input id="endDate" type="date"></label>
        <button id="downloadCsv" class="command-btn" type="button">导出当前回测 CSV</button>
        <button id="downloadParams" class="command-btn secondary" type="button">导出参数 JSON</button>
        <a class="command-btn secondary" href="../index.html" style="display:inline-flex;align-items:center;text-decoration:none;">返回汇总页</a>
      </div>
    </header>
    <main class="workspace">
      <aside class="data-panel">
        <div class="panel-heading"><h2>股票标的</h2><span id="dataCount" class="muted"></span></div>
        <div id="dataModules" class="module-list"></div>
      </aside>
      <section class="chart-panel">
        <div class="chart-header">
          <div><p id="chartEyebrow" class="eyebrow">Adjusted qfq_* price basis</p><h2 id="chartTitle">加载中</h2></div>
          <div id="statusPills" class="pill-row"></div>
        </div>
        <div id="warningBox" class="warning-box hidden"></div>
        <section class="summary-strip">
          <div><span class="summary-label">累计收益</span><strong id="cum">-</strong></div>
          <div><span class="summary-label">最大回撤</span><strong id="dd">-</strong></div>
          <div><span class="summary-label">Sharpe</span><strong id="sharpe">-</strong></div>
          <div><span class="summary-label">买入持有</span><strong id="bench">-</strong></div>
          <div><span class="summary-label">交易次数</span><strong id="trades">-</strong></div>
        </section>
        <div class="chart-grid">
          <section class="chart-block large"><div class="chart-block-head"><h3>复权价格、高低通道与交易信号</h3><span id="priceLabel" class="muted"></span></div><svg id="priceSvg" role="img" aria-label="价格通道图"></svg></section>
          <section class="chart-block"><div class="chart-block-head"><h3>策略净值、基准与回撤</h3><span id="equityLabel" class="muted"></span></div><svg id="equitySvg" role="img" aria-label="净值回撤图"></svg></section>
          <section class="chart-block"><div class="chart-block-head"><h3>参数敏感性</h3><span class="muted">当前标的 · 4 组通道参数</span></div><svg id="sensSvg" role="img" aria-label="参数敏感性图"></svg></section>
        </div>
        <section class="table-panel">
          <table><thead><tr><th>标的</th><th>累计收益</th><th>最大回撤</th><th>Sharpe</th><th>买入持有</th><th>交易次数</th><th>持仓占比</th></tr></thead><tbody id="metricsBody"></tbody></table>
        </section>
      </section>
      <aside class="param-panel">
        <div class="panel-heading"><h2>策略参数</h2><span id="paramState" class="muted">20/10</span></div>
        <section class="control-section"><h3>参数预设</h3><div class="preset-grid"><button type="button" data-preset="20,10">20/10</button><button type="button" data-preset="30,15">30/15</button><button type="button" data-preset="40,20">40/20</button><button type="button" data-preset="55,20">55/20</button></div></section>
        <section class="control-section"><h3>通道与 N 值</h3><div id="strategyControls"></div></section>
        <section class="control-section"><h3>资金与交易成本</h3><div id="moneyControls"></div></section>
        <section class="control-section"><h3>图表显示</h3><div class="toggle-grid"><label><input id="showEntry" type="checkbox" checked> 入场通道</label><label><input id="showExit" type="checkbox" checked> 退出通道</label><label><input id="showSignals" type="checkbox" checked> 买卖信号</label><label><input id="showBenchmark" type="checkbox" checked> 买入持有</label><label><input id="showDrawdown" type="checkbox" checked> 回撤</label></div></section>
        <section class="control-section"><h3>当前规则</h3><p id="ruleText" class="small-copy"></p></section>
      </aside>
    </main>
  </div>
  <script>
    const payload = __PAYLOAD__;
    const pct = v => v === null || Number.isNaN(v) ? "-" : (v * 100).toFixed(2) + "%";
    const num = v => v === null || Number.isNaN(v) ? "-" : Number(v).toFixed(2);
    const fmtMoney = v => v === null || Number.isNaN(v) ? "-" : Math.round(v).toLocaleString("zh-CN");
    const state = { datasetIndex: 0, entryWindow: 20, exitWindow: 10, nWindow: 20, stopN: 2, initialCash: 100000, riskPerUnit: 0.01, feeRate: 0.001, slippageRate: 0.0005, allowFractional: true };
    const controls = [["entryWindow","入场通道",2,80,1,"strategyControls"],["exitWindow","退出通道",2,60,1,"strategyControls"],["nWindow","N 值周期",2,60,1,"strategyControls"],["stopN","止损 N 倍数",0.5,5,0.1,"strategyControls"],["initialCash","初始资金",10000,1000000,10000,"moneyControls"],["riskPerUnit","单 unit 风险",0.002,0.05,0.001,"moneyControls"],["feeRate","手续费率",0,0.01,0.0001,"moneyControls"],["slippageRate","滑点率",0,0.01,0.0001,"moneyControls"]];
    function buildControls(){ controls.forEach(([key,label,min,max,step,target])=>{ const wrap=document.createElement("label"); wrap.className="control-row"; wrap.innerHTML=`<span>${label}</span><input id="${key}" type="number" min="${min}" max="${max}" step="${step}" value="${state[key]}">`; document.querySelector("#"+target).appendChild(wrap); wrap.querySelector("input").addEventListener("input", e=>{ const value=Number(e.target.value); if(!Number.isNaN(value)){ state[key]=value; render(); }}); }); }
    function initDates(){ const dates=payload.datasets.flatMap(d=>d.rows.map(r=>r.date)); document.querySelector("#startDate").value=dates.reduce((a,b)=>a<b?a:b); document.querySelector("#endDate").value=dates.reduce((a,b)=>a>b?a:b); document.querySelector("#startDate").addEventListener("input", render); document.querySelector("#endDate").addEventListener("input", render); }
    function initDataCards(){ const list=document.querySelector("#dataModules"); document.querySelector("#dataCount").textContent=`${payload.datasets.length} 个标的`; list.innerHTML=""; payload.datasets.forEach((d,i)=>{ const btn=document.createElement("button"); btn.type="button"; btn.className="data-card"; btn.dataset.index=i; btn.innerHTML=`<strong>${d.label}</strong><span>${d.market} · ${d.currency}</span><span>${d.startDate} 至 ${d.endDate}</span>`; btn.addEventListener("click",()=>{ state.datasetIndex=i; render(); }); list.appendChild(btn); }); }
    function initPresets(){ document.querySelectorAll("[data-preset]").forEach(btn=>{ btn.addEventListener("click",()=>{ const [entry, exit]=btn.dataset.preset.split(",").map(Number); state.entryWindow=entry; state.exitWindow=exit; document.querySelector("#entryWindow").value=entry; document.querySelector("#exitWindow").value=exit; render(); }); }); }
    function windowMax(rows, idx, key, win){ if(idx<win) return null; let out=-Infinity; for(let i=idx-win;i<idx;i+=1) out=Math.max(out, rows[i][key]); return out; }
    function windowMin(rows, idx, key, win){ if(idx<win) return null; let out=Infinity; for(let i=idx-win;i<idx;i+=1) out=Math.min(out, rows[i][key]); return out; }
    function enrichRows(baseRows,cfg){ const rows=baseRows.map(r=>({...r})); for(let i=0;i<rows.length;i+=1){ const r=rows[i]; r.tr=Math.max(r.high-r.low, Math.abs(r.high-r.preClose), Math.abs(r.low-r.preClose)); r.entryHigh=windowMax(rows,i,"high",cfg.entryWindow); r.exitLow=windowMin(rows,i,"low",cfg.exitWindow); r.n=null; } if(rows.length>=cfg.nWindow){ let sum=0; for(let i=0;i<cfg.nWindow;i+=1) sum+=rows[i].tr; rows[cfg.nWindow-1].n=sum/cfg.nWindow; for(let i=cfg.nWindow;i<rows.length;i+=1) rows[i].n=((cfg.nWindow-1)*rows[i-1].n+rows[i].tr)/cfg.nWindow; } for(let i=0;i<rows.length;i+=1){ rows[i].nForSignal=i>0?rows[i-1].n:null; rows[i].breakoutLong=rows[i].entryHigh!==null&&rows[i].nForSignal!==null&&rows[i].close>rows[i].entryHigh; rows[i].exitLong=rows[i].exitLow!==null&&rows[i].close<rows[i].exitLow; } return rows; }
    function backtest(baseRows,cfg){ const rows=enrichRows(baseRows,cfg); let cash=cfg.initialCash, shares=0, stopPrice=null, pending=null, runningMax=cfg.initialCash; for(let i=0;i<rows.length;i+=1){ const r=rows[i]; r.tradeAction="HOLD"; r.tradeReason=""; if(pending){ if(pending.side==="BUY"&&shares<=0&&cash>0&&pending.nForSize>0){ const executionPrice=r.open*(1+cfg.slippageRate); const unitShares=(cash*cfg.riskPerUnit)/(cfg.stopN*pending.nForSize); const maxAffordable=cash/(executionPrice*(1+cfg.feeRate)); let buyShares=Math.min(unitShares,maxAffordable); if(!cfg.allowFractional) buyShares=Math.floor(buyShares); if(buyShares>0){ const tradeValue=buyShares*executionPrice; const fee=tradeValue*cfg.feeRate; cash-=tradeValue+fee; shares=buyShares; stopPrice=executionPrice-cfg.stopN*pending.nForSize; r.tradeAction="BUY"; r.tradeReason=pending.reason; }} else if(pending.side==="SELL"&&shares>0){ const executionPrice=r.open*(1-cfg.slippageRate); const tradeValue=shares*executionPrice; const fee=tradeValue*cfg.feeRate; cash+=tradeValue-fee; shares=0; stopPrice=null; r.tradeAction="SELL"; r.tradeReason=pending.reason; } pending=null; } r.stopPrice=stopPrice; r.positionUnits=shares>0?1:0; r.shares=shares; r.cash=cash; r.portfolio=cash+shares*r.close; r.benchmark=cfg.initialCash*r.close/rows[0].close; runningMax=Math.max(runningMax,r.portfolio); r.drawdown=r.portfolio/runningMax-1; r.stopLong=shares>0&&stopPrice!==null&&r.close<=stopPrice; if(i<rows.length-1){ if(shares>0){ if(r.stopLong) pending={side:"SELL",reason:"2N stop loss"}; else if(r.exitLong) pending={side:"SELL",reason:`${cfg.exitWindow}-day low exit`}; } else if(r.breakoutLong){ pending={side:"BUY",reason:`${cfg.entryWindow}-day high breakout`,nForSize:r.nForSignal}; } } } for(let i=0;i<rows.length;i+=1){ rows[i].strategyReturn=i===0?0:rows[i].portfolio/rows[i-1].portfolio-1; rows[i].benchmarkReturn=i===0?0:rows[i].benchmark/rows[i-1].benchmark-1; } return rows; }
    function metrics(rows,cfg){ if(!rows.length) return {finalValue:cfg.initialCash,cumulative:0,maxDrawdown:0,sharpe:null,benchmarkCumulative:0,tradeCount:0,exposure:0}; const finalValue=rows.at(-1).portfolio; const cumulative=finalValue/cfg.initialCash-1; const mean=rows.reduce((s,r)=>s+r.strategyReturn,0)/Math.max(1,rows.length); const variance=rows.length>1?rows.reduce((s,r)=>s+Math.pow(r.strategyReturn-mean,2),0)/(rows.length-1):0; const std=Math.sqrt(variance); const sharpe=std>0?Math.sqrt(252)*mean/std:null; return {finalValue,cumulative,maxDrawdown:Math.min(...rows.map(r=>r.drawdown)),sharpe,benchmarkCumulative:rows.at(-1).benchmark/rows[0].benchmark-1,tradeCount:rows.filter(r=>r.tradeAction==="BUY"||r.tradeAction==="SELL").length,exposure:rows.filter(r=>r.positionUnits>0).length/Math.max(1,rows.length)}; }
    function currentConfig(){ return { entryWindow:Math.max(2,Math.round(state.entryWindow)), exitWindow:Math.max(2,Math.round(state.exitWindow)), nWindow:Math.max(2,Math.round(state.nWindow)), stopN:Math.max(.1,state.stopN), initialCash:Math.max(1,state.initialCash), riskPerUnit:Math.max(.0001,state.riskPerUnit), feeRate:Math.max(0,state.feeRate), slippageRate:Math.max(0,state.slippageRate), allowFractional:state.allowFractional }; }
    function filteredRows(dataset){ const start=document.querySelector("#startDate").value; const end=document.querySelector("#endDate").value; return dataset.rows.filter(r=>(!start||r.date>=start)&&(!end||r.date<=end)); }
    function points(rows,key,w,h,pad,minV,maxV){ return rows.map((r,i)=>{ const x=pad+i*(w-2*pad)/Math.max(1,rows.length-1); const v=r[key]; const y=h-pad-(v-minV)/Math.max(1e-9,maxV-minV)*(h-2*pad); return [x,y,v]; }).filter(p=>p[2]!==null&&!Number.isNaN(p[2])); }
    function pathFrom(pts){ return pts.map((p,i)=>(i?"L":"M")+p[0].toFixed(1)+","+p[1].toFixed(1)).join(" "); }
    function drawPrice(dataset,rows,cfg){ const svg=document.querySelector("#priceSvg"); const w=svg.clientWidth||760, h=420, pad=42; const visible=rows.slice(-190); const values=visible.flatMap(r=>[r.close,r.entryHigh,r.exitLow,r.stopPrice]).filter(v=>v!==null); const minV=Math.min(...values), maxV=Math.max(...values), span=maxV-minV||1, lo=minV-span*.08, hi=maxV+span*.08; const close=points(visible,"close",w,h,pad,lo,hi), entry=points(visible,"entryHigh",w,h,pad,lo,hi), exit=points(visible,"exitLow",w,h,pad,lo,hi), stop=points(visible,"stopPrice",w,h,pad,lo,hi); const marks=visible.map((r,i)=>({...r,x:pad+i*(w-2*pad)/Math.max(1,visible.length-1)})).filter(r=>r.tradeAction==="BUY"||r.tradeAction==="SELL"); svg.setAttribute("viewBox",`0 0 ${w} ${h}`); svg.innerHTML=`<rect x="${pad}" y="20" width="${w-2*pad}" height="${h-2*pad}" fill="white" stroke="#d8dee8"/>${document.querySelector("#showExit").checked?`<path d="${pathFrom(exit)}" fill="none" stroke="#0f766e" stroke-width="2"/>`:""}${document.querySelector("#showEntry").checked?`<path d="${pathFrom(entry)}" fill="none" stroke="#d9822b" stroke-width="2"/>`:""}<path d="${pathFrom(stop)}" fill="none" stroke="#b94d4d" stroke-width="1.5" stroke-dasharray="5 5"/><path d="${pathFrom(close)}" fill="none" stroke="#213547" stroke-width="2.5"/>${document.querySelector("#showSignals").checked?marks.map(m=>`<circle cx="${m.x}" cy="${points([m],"close",w,h,pad,lo,hi)[0]?.[1]??0}" r="5" fill="${m.tradeAction==="BUY"?"#2e8f57":"#b94d4d"}"><title>${m.date} ${m.tradeAction} ${m.tradeReason}</title></circle>`).join(""):""}<text x="${pad}" y="14" fill="#667385" font-size="12">${dataset.label} · ${cfg.entryWindow}/${cfg.exitWindow} · N=${cfg.nWindow}</text><text x="${pad}" y="${h-8}" fill="#64748b" font-size="12">${visible[0]?.date??""}</text><text x="${w-pad}" y="${h-8}" fill="#64748b" font-size="12" text-anchor="end">${visible.at(-1)?.date??""}</text>`; }
    function drawEquity(rows){ const svg=document.querySelector("#equitySvg"); const w=svg.clientWidth||760, h=300, pad=42; const values=rows.flatMap(r=>[r.portfolio,r.benchmark]).filter(v=>v!==null); const minV=Math.min(...values), maxV=Math.max(...values), span=maxV-minV||1; const port=points(rows,"portfolio",w,190,pad,minV-span*.08,maxV+span*.08); const bench=points(rows,"benchmark",w,190,pad,minV-span*.08,maxV+span*.08); const minDd=Math.min(...rows.map(r=>r.drawdown),-.01); const dd=rows.map((r,i)=>[pad+i*(w-2*pad)/Math.max(1,rows.length-1),260-(r.drawdown-minDd)/(0-minDd)*70]).filter(p=>Number.isFinite(p[1])); svg.setAttribute("viewBox",`0 0 ${w} ${h}`); svg.innerHTML=`<rect x="${pad}" y="20" width="${w-2*pad}" height="160" fill="white" stroke="#d8dee8"/>${document.querySelector("#showBenchmark").checked?`<path d="${pathFrom(bench)}" fill="none" stroke="#64748b" stroke-width="2"/>`:""}<path d="${pathFrom(port)}" fill="none" stroke="#1f6fb2" stroke-width="2.5"/><rect x="${pad}" y="205" width="${w-2*pad}" height="70" fill="white" stroke="#d8dee8"/>${document.querySelector("#showDrawdown").checked?`<path d="${pathFrom(dd)}" fill="none" stroke="#b94d4d" stroke-width="2"/>`:""}<text x="${pad}" y="294" fill="#64748b" font-size="12">蓝线为策略净值，灰线为买入持有，红线为回撤</text>`; }
    function drawSensitivity(dataset,cfg){ const svg=document.querySelector("#sensSvg"); const w=svg.clientWidth||760, h=300, pad=38; const combos=[[20,10],[30,15],[40,20],[55,20]]; const base=filteredRows(dataset); const rows=combos.map(([entryWindow,exitWindow])=>{ const localCfg={...cfg,entryWindow,exitWindow}; return {parameter:`${entryWindow}/${exitWindow}`,...metrics(backtest(base,localCfg),localCfg)}; }); const cell=(w-2*pad)/rows.length; svg.setAttribute("viewBox",`0 0 ${w} ${h}`); svg.innerHTML=rows.map((r,i)=>{ const x=pad+i*cell+6; const fill=r.cumulative>=0?"#bfe6ce":"#f2c4c4"; return `<rect x="${x}" y="38" width="${cell-12}" height="175" rx="8" fill="${fill}" stroke="#d8dee8"/><text x="${x+(cell-12)/2}" y="76" text-anchor="middle" font-size="17" fill="#172033">${r.parameter}</text><text x="${x+(cell-12)/2}" y="116" text-anchor="middle" font-size="14" fill="#172033">收益 ${pct(r.cumulative)}</text><text x="${x+(cell-12)/2}" y="148" text-anchor="middle" font-size="14" fill="#172033">回撤 ${pct(r.maxDrawdown)}</text><text x="${x+(cell-12)/2}" y="180" text-anchor="middle" font-size="14" fill="#172033">交易 ${r.tradeCount}</text>`; }).join(""); }
    function updateMetricsTable(cfg){ document.querySelector("#metricsBody").innerHTML=payload.datasets.map(d=>{ const result=backtest(filteredRows(d),cfg); const s=metrics(result,cfg); return `<tr><td>${d.label}</td><td>${pct(s.cumulative)}</td><td>${pct(s.maxDrawdown)}</td><td>${num(s.sharpe)}</td><td>${pct(s.benchmarkCumulative)}</td><td>${s.tradeCount}</td><td>${pct(s.exposure)}</td></tr>`; }).join(""); }
    function render(){ const cfg=currentConfig(); const dataset=payload.datasets[state.datasetIndex]; const baseRows=filteredRows(dataset); const warning=document.querySelector("#warningBox"); document.querySelectorAll(".data-card").forEach((card,idx)=>card.classList.toggle("active",idx===state.datasetIndex)); document.querySelectorAll("[data-preset]").forEach(btn=>btn.classList.toggle("active",btn.dataset.preset===`${cfg.entryWindow},${cfg.exitWindow}`)); document.querySelector("#paramState").textContent=`${cfg.entryWindow}/${cfg.exitWindow}`; document.querySelector("#chartTitle").textContent=dataset.label; document.querySelector("#priceLabel").textContent=`入场 ${cfg.entryWindow} · 退出 ${cfg.exitWindow} · N ${cfg.nWindow}`; document.querySelector("#equityLabel").textContent=`初始资金 ${fmtMoney(cfg.initialCash)}`; document.querySelector("#statusPills").innerHTML=[dataset.market,dataset.currency,`行数 ${baseRows.length}`,`风险 ${pct(cfg.riskPerUnit)}`].map(x=>`<span class="pill">${x}</span>`).join(""); document.querySelector("#ruleText").textContent=`收盘价突破前一日 ${cfg.entryWindow} 日高点通道后，下一交易日开盘买入；跌破 ${cfg.exitWindow} 日低点通道或触发 ${cfg.stopN}N 止损后，下一交易日开盘卖出。仓位按单 unit 风险 ${pct(cfg.riskPerUnit)} 计算。`; if(baseRows.length<=Math.max(cfg.entryWindow,cfg.exitWindow,cfg.nWindow)+2){ warning.textContent="当前日期区间过短，可能不足以完成通道和 N 值预热。"; warning.classList.remove("hidden"); } else warning.classList.add("hidden"); const result=backtest(baseRows,cfg); const s=metrics(result,cfg); document.querySelector("#cum").textContent=pct(s.cumulative); document.querySelector("#dd").textContent=pct(s.maxDrawdown); document.querySelector("#sharpe").textContent=num(s.sharpe); document.querySelector("#bench").textContent=pct(s.benchmarkCumulative); document.querySelector("#trades").textContent=s.tradeCount; drawPrice(dataset,result,cfg); drawEquity(result); drawSensitivity(dataset,cfg); updateMetricsTable(cfg); window.currentBacktestRows=result; window.currentBacktestConfig=cfg; }
    function download(name,text,type="text/plain"){ const blob=new Blob([text],{type}); const url=URL.createObjectURL(blob); const a=document.createElement("a"); a.href=url; a.download=name; a.click(); URL.revokeObjectURL(url); }
    document.querySelector("#downloadCsv").addEventListener("click",()=>{ const rows=window.currentBacktestRows||[]; const headers=["date","open","high","low","close","entryHigh","exitLow","n","stopPrice","tradeAction","portfolio","benchmark","drawdown"]; const csv=[headers.join(","),...rows.map(r=>headers.map(h=>r[h]??"").join(","))].join("\n"); download("task4_turtle_current_backtest.csv",csv,"text/csv;charset=utf-8"); });
    document.querySelector("#downloadParams").addEventListener("click",()=>download("task4_turtle_params.json",JSON.stringify(window.currentBacktestConfig||currentConfig(),null,2),"application/json"));
    document.querySelectorAll(".toggle-grid input").forEach(input=>input.addEventListener("change",render));
    window.addEventListener("resize",render);
    buildControls(); initDates(); initDataCards(); initPresets(); render();
  </script>
</body>
</html>
"""
    HTML_PATH.write_text(template.replace("__PAYLOAD__", payload), encoding="utf-8")
    (TASK_DIR / "index.html").write_text(
        """<!doctype html>
<html lang="zh-CN">
<head><meta charset="utf-8"><meta name="viewport" content="width=device-width, initial-scale=1"><meta http-equiv="refresh" content="0; url=turtle_strategy_report.html"><title>Task 4 海龟策略回测</title></head>
<body><p><a href="turtle_strategy_report.html">打开 Task 4 海龟策略回测报告</a></p></body>
</html>
""",
        encoding="utf-8",
    )


def build_notebook() -> None:
    cells = [
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "# Task 4 海龟策略回测\\n",
                "\\n",
                "本 Notebook 使用复权后的 `qfq_*` 价格字段计算海龟策略。通道使用 `shift(1)`，信号采用“收盘确认、下一交易日开盘执行”，避免未来函数。\\n",
            ],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 1. 策略思想\\n",
                "\\n",
                "海龟策略的核心思想是用高低点通道识别趋势突破，用 ATR/N 衡量波动，用 2N 止损和风险定仓控制单笔风险。它的优势在于规则透明、纪律性强、能够抓住持续趋势，同时也会在震荡市场中产生假突破。\\n",
            ],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "from pathlib import Path\\n",
                "import pandas as pd\\n",
                "from turtle_strategy import TurtleStrategyConfig, add_turtle_indicators\\n",
                "from backtester import TurtleBacktestConfig, run_turtle_backtest, calculate_performance_metrics\\n",
                "\\n",
                "PROJECT_ROOT = Path('..').resolve()\\n",
                "DATA_PATH = PROJECT_ROOT / 'data/processed/equities/daily/smic_h_00981_HK_daily_20250703_20260703.csv'\\n",
                "strategy_config = TurtleStrategyConfig(entry_window=20, exit_window=10, n_window=20, stop_n=2.0)\\n",
                "backtest_config = TurtleBacktestConfig(initial_cash=100000, risk_per_unit=0.01, fee_rate=0.001, slippage_rate=0.0005)\\n",
            ],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "data = pd.read_csv(DATA_PATH, dtype={'trade_date': str, 'ts_code': str}, encoding='utf-8-sig')\\n",
                "data['date'] = pd.to_datetime(data['trade_date'], format='%Y%m%d')\\n",
                "data = data.sort_values('trade_date').reset_index(drop=True)\\n",
                "data.head()\\n",
            ],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 2. 计算高低点通道与 ATR/N\\n",
                "\\n",
                "20 日高点通道用于入场，10 日低点通道用于退出。真实波幅 TR 取 `high-low`、`abs(high-pre_close)`、`abs(low-pre_close)` 的最大值，N 值使用海龟原版 20 日平滑。\\n",
            ],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "signals = add_turtle_indicators(data, strategy_config)\\n",
                "signals[['trade_date','qfq_close','tr','n_20','entry_high_20','exit_low_10','breakout_long','exit_long']].tail(10)\\n",
            ],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 3. 交易信号与模拟回测\\n",
                "\\n",
                "当收盘价突破前一日已知的 20 日高点通道且当前无仓位时，下一交易日开盘买入；当收盘价跌破 10 日低点通道或触发 2N 止损时，下一交易日开盘卖出。\\n",
            ],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "portfolio = run_turtle_backtest(data, strategy_config, backtest_config)\\n",
                "metrics = calculate_performance_metrics(portfolio, backtest_config)\\n",
                "metrics\\n",
            ],
        },
        {
            "cell_type": "markdown",
            "metadata": {},
            "source": [
                "## 4. 可视化和参数敏感性\\n",
                "\\n",
                "完整图表、批量回测结果和参数敏感性结果已由 `build_task4_deliverables.py` 生成到 `outputs/` 目录，并同步写入 HTML 和 Word 报告。\\n",
            ],
        },
        {
            "cell_type": "code",
            "execution_count": None,
            "metadata": {},
            "outputs": [],
            "source": [
                "summary = pd.read_csv('outputs/turtle_core_metrics_summary.csv')\\n",
                "summary[['label','cumulative_return','max_drawdown','sharpe_ratio','benchmark_cumulative_return','trade_count']]\\n",
            ],
        },
    ]
    for idx, cell in enumerate(cells, start=1):
        cell["id"] = f"task4-cell-{idx:02d}"
    notebook = {
        "cells": cells,
        "metadata": {
            "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"},
            "language_info": {"name": "python", "pygments_lexer": "ipython3"},
        },
        "nbformat": 4,
        "nbformat_minor": 5,
    }
    NOTEBOOK_PATH.write_text(json.dumps(notebook, ensure_ascii=False, indent=1), encoding="utf-8")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths_inches: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=10.5, bold=False, color: str | None = None) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=False) -> None:
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.5
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.first_line_indent = Cm(0.74) if first_line else Cm(0)
    for run in paragraph.runs:
        if not run.font.size:
            set_run_font(run)


def add_para(doc: Document, text: str = "", bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_line=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, bold=bold)
    format_paragraph(p, align=align, first_line=first_line)
    return p


def add_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, bold=True)
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, first_line=False)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, bold=True)
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)


def style_table(table, widths_inches: list[float] | None = None) -> None:
    table.style = "Table Grid"
    if widths_inches:
        set_table_width(table, widths_inches)
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:tblHeader")) is None:
        tr_pr.append(OxmlElement("w:tblHeader"))
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                format_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
                for run in paragraph.runs:
                    set_run_font(run, size=9.5, bold=row_idx == 0)
            if row_idx == 0:
                set_cell_shading(cell, "D9EAF7")


def add_table(doc: Document, caption: str, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = str(header)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    style_table(table, widths)


def add_figure(doc: Document, path: Path, caption: str, interpretation: str, width=5.75) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    shape = run.add_picture(str(path), width=Inches(width))
    shape._inline.docPr.set("title", caption)
    shape._inline.docPr.set("descr", interpretation[:220])
    format_paragraph(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_caption(doc, caption)
    add_para(doc, interpretation)


def build_word_doc(metrics_rows: list[dict[str, object]], sens: pd.DataFrame) -> None:
    TASK4_DOC_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("TASK4 海龟策略实现与回测分析报告")
    set_run_font(run, size=10.5, bold=True)
    format_paragraph(title, align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "姓名：程冰晖", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "策略：Turtle Core 20/10，20 日 N 值，2N 止损", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "数据口径：复权日线价格 qfq_*", align=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
    add_para(doc, "")

    add_heading(doc, "一、海龟策略的核心思想与优势")
    add_para(
        doc,
        "海龟策略是一类典型的趋势跟踪策略。它不试图预测价格顶部或底部，而是用价格突破近期高点来确认趋势可能启动，"
        "再用近期低点、平均真实波幅和固定风险预算来管理退出与止损。本报告采用核心版规则：当复权收盘价突破前一日已知的 20 日高点通道时，"
        "下一交易日开盘买入；当复权收盘价跌破 10 日低点通道或触发 2N 止损时，下一交易日开盘卖出。"
    )
    add_para(
        doc,
        "该策略的关键优势包括：第一，规则明确，交易条件可以被程序完整复现；第二，顺势而为，能够在强趋势阶段保留头寸；"
        "第三，使用 ATR/N 值衡量波动并按风险定仓，使不同价格和波动水平的股票具有相对一致的单笔风险；第四，止损条件事前定义，"
        "有助于减少主观判断带来的执行偏差。其局限也很明显：在横盘震荡中容易遇到假突破，样本较短时回测结果稳定性有限。"
    )

    add_heading(doc, "二、关键概念说明")
    add_table(
        doc,
        "表1 海龟策略关键概念",
        ["概念", "本报告中的计算方式", "作用"],
        [
            ["高点通道", "entry_high_20 = rolling_max(qfq_high, 20).shift(1)", "用于识别向上突破入场信号"],
            ["低点通道", "exit_low_10 = rolling_min(qfq_low, 10).shift(1)", "用于识别趋势转弱后的退出信号"],
            ["真实波幅 TR", "max(high-low, abs(high-pre_close), abs(low-pre_close))", "度量单日真实波动范围"],
            ["N 值", "20 日 TR 的海龟平滑值", "作为波动率单位，用于定仓和止损"],
            ["2N 止损", "stop_price = entry_price - 2 * entry_n", "限制单笔交易风险"],
        ],
        [1.25, 3.15, 2.1],
    )
    add_para(doc, "表1中所有价格字段均使用复权后的 qfq_* 字段。通道计算使用 shift(1)，即第 t 日判断信号时，只使用第 t-1 日及以前的高低点区间，避免未来函数。")

    add_heading(doc, "三、Python 编程实现")
    add_para(
        doc,
        "程序首先从 data/processed/equities/daily/ 读取已存储的股票日线 CSV，并检查 trade_date 是否唯一、ts_code 是否单一、"
        "qfq_open、qfq_high、qfq_low、qfq_close 和 qfq_pre_close 是否完整。随后计算 TR、20 日 N 值、20 日高点通道和 10 日低点通道。"
        "交易信号采用收盘确认、下一交易日开盘执行的方式，买入执行价为 next_qfq_open 叠加滑点，卖出执行价为 next_qfq_open 扣减滑点。"
    )
    add_para(
        doc,
        "仓位管理使用风险定仓：unit_risk_cash = current_equity × 1%，unit_shares = unit_risk_cash / (2 × N)。"
        "这一设置意味着如果入场后价格向下移动 2N，理论亏损约为账户权益的 1%。回测同时扣除 0.1% 手续费和 0.05% 滑点，"
        "并计算策略净值、买入持有基准、累计收益、最大回撤、夏普比率、持仓天数占比和交易次数等指标。"
    )

    add_figure(
        doc,
        TASK4_DOC_FIG_DIR / "fig1_turtle_signals.png",
        "图1 中芯国际港股海龟策略通道与交易信号",
        "图1展示了复权收盘价与 20 日入场通道、10 日退出通道的相对位置。绿色三角为买入执行点，红色三角为卖出执行点。可以看到，海龟策略只在价格突破近期高点后入场，并在趋势回落或触发止损后离场，信号数量较少，体现出趋势跟踪策略等待突破的特点。",
    )
    add_figure(
        doc,
        TASK4_DOC_FIG_DIR / "fig2_equity_drawdown.png",
        "图2 策略净值、买入持有基准与回撤",
        "图2比较了海龟策略净值和买入持有基准，并在下方展示回撤。若策略净值在趋势阶段上行但回撤被控制，说明突破规则和 2N 止损在样本内起到了一定风险约束作用；若策略长期落后于买入持有，则说明该样本可能更适合买入持有或参数需要调整。",
    )

    add_heading(doc, "四、多标的回测结果")
    metric_table = []
    for row in metrics_rows:
        label_for_table = (
            str(row["label"])
            .replace("（", "\n")
            .replace("）", "")
            .replace("中芯国际", "中芯")
            .replace("长江电力", "长电")
            .replace("比亚迪", "比亚迪")
        )
        metric_table.append(
            [
                label_for_table,
                pct(float(row["cumulative_return"])),
                pct(float(row["max_drawdown"])),
                num(float(row["sharpe_ratio"])),
                pct(float(row["benchmark_cumulative_return"])),
                str(int(row["trade_count"])),
            ]
        )
    add_table(
        doc,
        "表2 多标的海龟策略绩效汇总",
        ["标的", "累计收益", "最大回撤", "Sharpe", "买入持有", "交易次数"],
        metric_table,
        [1.75, 0.95, 0.95, 0.85, 0.95, 0.75],
    )
    add_para(doc, "表2显示，同一套规则在不同股票上的表现差异明显。海龟策略依赖持续趋势，因此趋势强、波动扩张后能延续的标的更容易产生较好收益；如果价格频繁突破后回落，则交易次数可能增加，但收益和夏普比率未必改善。")
    add_figure(
        doc,
        TASK4_DOC_FIG_DIR / "fig3_multi_stock_metrics.png",
        "图3 多标的海龟策略绩效对比",
        "图3把累计收益、最大回撤和 Sharpe 放在同一张对比图中。最大回撤以绝对值展示，便于观察风险大小。该图说明，海龟策略不是对所有股票都稳定有效，策略表现与股票自身趋势结构、波动水平和样本期行情密切相关。",
    )

    add_heading(doc, "五、参数调节与适应场景")
    sens_rows = [
        [
            str(row["parameter"]),
            pct(float(row["cumulative_return"])),
            pct(float(row["max_drawdown"])),
            num(float(row["sharpe_ratio"])),
            str(int(row["trade_count"])),
        ]
        for _, row in sens.iterrows()
    ]
    add_table(
        doc,
        "表3 中芯国际港股参数敏感性",
        ["入场/退出", "累计收益", "最大回撤", "Sharpe", "交易次数"],
        sens_rows,
        [1.2, 1.1, 1.1, 1.0, 1.0],
    )
    add_figure(
        doc,
        TASK4_DOC_FIG_DIR / "fig4_parameter_sensitivity.png",
        "图4 核心参数敏感性分析",
        "图4展示不同通道周期下的收益、回撤、夏普比率和交易次数。短周期如 20/10 对价格变化更敏感，可能更早捕捉突破，但也更容易受到震荡噪声影响；长周期如 55/20 信号更少，通常要求趋势更强、更持久。",
    )
    add_para(
        doc,
        "综合本次实验，海龟法则更适合趋势持续性较强、突破后有延续空间的行情。对于横盘震荡、跳空频繁或趋势很短的股票，"
        "该策略可能反复触发买入和止损，导致交易成本增加。实际应用时应结合更长样本、交易单位、流动性、涨跌停和组合级风险上限进行进一步检验。"
    )

    add_heading(doc, "六、结论")
    add_para(
        doc,
        "本次 Task4 完成了数据加载、通道计算、ATR/N 值计算、买卖信号生成、可视化和模拟回测。"
        "结果表明，海龟策略规则清晰、风险控制明确，适合捕捉持续趋势；但它对震荡行情较敏感，且一年左右样本无法证明长期稳定性。"
        "本报告仅用于课程学习和历史回测分析，不构成投资建议。"
    )

    doc.save(DOCX_PATH)


def build_summary_page() -> None:
    task1_path = PROJECT_ROOT / "task1_analysis.html"
    original_index = PROJECT_ROOT / "index.html"
    if original_index.exists() and not task1_path.exists():
        task1_path.write_bytes(original_index.read_bytes())

    cards = [
        ("Task 1", "中际旭创股价分析", "复权价格、K 线、回撤和波动率分析。", "task1_analysis.html"),
        ("Task 2", "Indicator Lab", "RSI、MACD、布林带、ATR、MA、KDJ、OBV 交互式指标实验台。", "task2_indicator_lab/index.html"),
        ("Task 3", "双均线策略回测", "5/15 日均线信号、模拟交易、绩效对比和参数敏感性。", "task3_double_ma_strategy/index.html"),
        ("Task 4", "海龟策略回测", "唐奇安通道、ATR/N 值、2N 止损、风险定仓和多标的回测。", "task4_turtle_strategy/index.html"),
    ]
    card_html = "\n".join(
        f"""<a class="task-card" href="{href}">
        <span>{label}</span>
        <h2>{title}</h2>
        <p>{desc}</p>
      </a>"""
        for label, title, desc, href in cards
    )
    original_index.write_text(
        f"""<!doctype html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1">
  <title>AI Quant Study · Task 汇总</title>
  <style>
    :root {{ --ink:#172033; --muted:#64748b; --line:#d8dee8; --bg:#f6f8fb; --panel:#ffffff; --blue:#1d4ed8; }}
    * {{ box-sizing:border-box; }}
    body {{ margin:0; font-family:-apple-system,BlinkMacSystemFont,"Segoe UI","PingFang SC","Microsoft YaHei",sans-serif; color:var(--ink); background:var(--bg); }}
    header {{ min-height:34vh; padding:42px 36px 30px; background:#fff; border-bottom:1px solid var(--line); display:flex; align-items:flex-end; }}
    .hero {{ width:min(1120px, calc(100vw - 36px)); margin:0 auto; }}
    .eyebrow {{ color:var(--blue); font-weight:700; letter-spacing:.08em; text-transform:uppercase; margin:0 0 10px; }}
    h1 {{ margin:0 0 12px; font-size:clamp(34px,5vw,58px); letter-spacing:0; }}
    p {{ margin:0; color:var(--muted); line-height:1.7; }}
    main {{ width:min(1120px, calc(100vw - 36px)); margin:24px auto 48px; }}
    .grid {{ display:grid; grid-template-columns:repeat(2,minmax(0,1fr)); gap:18px; }}
    .task-card {{ display:block; min-height:190px; padding:22px; border:1px solid var(--line); border-radius:8px; background:var(--panel); color:inherit; text-decoration:none; transition:transform .16s ease, box-shadow .16s ease; }}
    .task-card:hover {{ transform:translateY(-2px); box-shadow:0 12px 28px rgba(23,32,51,.08); }}
    .task-card span {{ color:var(--blue); font-weight:700; }}
    .task-card h2 {{ margin:18px 0 10px; font-size:24px; letter-spacing:0; }}
    @media (max-width:760px) {{ .grid {{ grid-template-columns:1fr; }} header {{ min-height:26vh; padding:30px 18px; }} }}
  </style>
</head>
<body>
  <header>
    <section class="hero">
      <p class="eyebrow">AI Quant Study</p>
      <h1>四个任务汇总入口</h1>
      <p>这里集中访问 Task 1 到 Task 4 的 HTML 成果页。所有分析均使用课程数据和复权价格口径，仅用于学习研究，不构成投资建议。</p>
    </section>
  </header>
  <main>
    <section class="grid">
      {card_html}
    </section>
  </main>
</body>
</html>
""",
        encoding="utf-8",
    )


def main() -> None:
    TASK_DIR.mkdir(parents=True, exist_ok=True)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    portfolios, metrics_rows = build_backtests()
    sens = parameter_sensitivity()
    build_figures(portfolios, metrics_rows, sens)
    build_interactive_html_report()
    build_notebook()
    build_word_doc(metrics_rows, sens)
    build_summary_page()
    print(f"Wrote {DOCX_PATH}")
    print(f"Wrote {HTML_PATH}")
    print(f"Wrote {NOTEBOOK_PATH}")
    print(f"Wrote {PROJECT_ROOT / 'index.html'}")


if __name__ == "__main__":
    main()
