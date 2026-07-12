from __future__ import annotations

import json
from pathlib import Path

import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from matplotlib import font_manager
from sklearn.metrics import confusion_matrix, roc_auc_score, roc_curve
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
TASK = ROOT / "task5_machine_learning"
OUTPUT = TASK / "outputs"
FIG = OUTPUT / "figures"
DOC_DIR = ROOT / "TASK5"
DOC_FIG = DOC_DIR / "figures"
DOCX = DOC_DIR / "程冰晖 TASK5.docx"
HTML = TASK / "machine_learning_report.html"

for d in [FIG, DOC_DIR, DOC_FIG]: d.mkdir(parents=True, exist_ok=True)

for name in ["Hiragino Sans GB", "Heiti SC", "Arial Unicode MS", "STSong", "SimHei"]:
    if name in {f.name for f in font_manager.fontManager.ttflist}:
        plt.rcParams["font.sans-serif"] = [name, "DejaVu Sans"]
        break
plt.rcParams["axes.unicode_minus"] = False

METRICS = pd.read_csv(OUTPUT / "metrics_all.csv")
PRED = pd.read_csv(OUTPUT / "predictions_all.csv", parse_dates=["date"])
SPLITS = pd.read_csv(OUTPUT / "split_summary.csv", parse_dates=["start", "end"])
DATA = pd.read_csv(ROOT / "data/processed/equities/daily/ml_equities_daily_20180101_20260712.csv", dtype={"trade_date": str})

LABELS = {"00981.HK": "中芯国际H股", "01211.HK": "比亚迪H股", "688981.SH": "中芯国际A股"}
MODEL_CN = {
    "logistic_regression": "逻辑回归", "decision_tree_classifier": "决策树",
    "random_forest_classifier": "随机森林", "knn_classifier": "KNN",
    "linear_regression": "线性回归", "decision_tree_regressor": "决策树回归",
    "random_forest_regressor": "随机森林回归", "knn_regressor": "KNN回归",
}

def pct(v): return "-" if pd.isna(v) else f"{v*100:.2f}%"
def num(v, n=3): return "-" if pd.isna(v) else f"{v:.{n}f}"

def make_figures():
    # Figure 1: data split timeline
    fig, ax = plt.subplots(figsize=(11.5, 4.8))
    colors = {"train": "#1f6fb2", "validation": "#d9822b", "test": "#2e8f57"}
    y = 0
    yt, yl = [], []
    for symbol in LABELS:
        for ds in ["train", "validation", "test"]:
            r = SPLITS[(SPLITS.symbol == symbol) & (SPLITS.dataset == ds)].iloc[0]
            ax.plot([r.start, r.end], [y, y], lw=12, color=colors[ds], solid_capstyle="butt")
            ax.text(r.start, y + .2, f"{ds} {int(r.rows)}", fontsize=9, color=colors[ds])
        yt.append(y); yl.append(f"{LABELS[symbol]}\n{symbol}"); y += 1
    ax.set_yticks(yt); ax.set_yticklabels(yl); ax.set_title("图1 三只股票训练集、验证集与测试集时间划分")
    ax.set_xlabel("交易日期"); ax.grid(axis="x", alpha=.25); fig.tight_layout()
    fig.savefig(FIG / "fig1_time_split.png", dpi=180); fig.savefig(DOC_FIG / "fig1_time_split.png", dpi=180); plt.close(fig)

    # Figure 2: Test AUC
    auc = METRICS[(METRICS.dataset == "test") & (METRICS.metric == "AUC")].copy()
    auc["股票"] = auc.symbol.map(LABELS); auc["模型"] = auc.model.map(MODEL_CN)
    pivot = auc.pivot(index="股票", columns="模型", values="value").reindex(columns=["逻辑回归","决策树","随机森林","KNN"])
    ax = pivot.plot(kind="bar", figsize=(11.5, 5.6), color=["#1f6fb2","#d9822b","#2e8f57","#7c5ab8"])
    ax.axhline(.5, color="#667385", ls="--", lw=1); ax.set_ylim(0, 1); ax.set_ylabel("AUC")
    ax.set_title("图2 测试集分类模型 AUC 对比"); ax.tick_params(axis="x", rotation=0); ax.legend(title="模型", ncol=4, loc="upper center")
    for c in ax.containers: ax.bar_label(c, fmt="%.3f", fontsize=8, padding=2)
    ax.figure.tight_layout(); ax.figure.savefig(FIG / "fig2_auc_comparison.png", dpi=180); ax.figure.savefig(DOC_FIG / "fig2_auc_comparison.png", dpi=180); plt.close(ax.figure)

    # Figure 3: ROC for random forest, all symbols
    fig, ax = plt.subplots(figsize=(8.5, 6.5))
    q = PRED[(PRED.dataset == "test") & (PRED.task_type == "classification") & (PRED.model == "random_forest_classifier")]
    for symbol, g in q.groupby("symbol"):
        fpr, tpr, _ = roc_curve(g.actual, g.probability); aucv = roc_auc_score(g.actual, g.probability)
        ax.plot(fpr, tpr, lw=2.2, label=f"{LABELS[symbol]} AUC={aucv:.3f}")
    ax.plot([0,1],[0,1],"--",color="#667385"); ax.set_xlabel("假正率 FPR"); ax.set_ylabel("真正率 TPR")
    ax.set_title("图3 随机森林测试集 ROC 曲线"); ax.legend(); ax.grid(alpha=.25); fig.tight_layout()
    fig.savefig(FIG / "fig3_roc_random_forest.png", dpi=180); fig.savefig(DOC_FIG / "fig3_roc_random_forest.png", dpi=180); plt.close(fig)

    # Figure 4: confusion matrices for RF
    fig, axes = plt.subplots(1, 3, figsize=(11.5, 3.7))
    for ax, (symbol, g) in zip(axes, q.groupby("symbol")):
        cm = confusion_matrix(g.actual, g.prediction, labels=[0,1])
        im=ax.imshow(cm,cmap="Blues")
        for i in range(2):
            for j in range(2): ax.text(j,i,int(cm[i,j]),ha="center",va="center",fontsize=13)
        ax.set_xticks([0,1]); ax.set_yticks([0,1]); ax.set_xlabel("预测类别"); ax.set_ylabel("真实类别"); ax.set_title(LABELS[symbol])
    fig.suptitle("图4 随机森林测试集混淆矩阵", fontsize=14); fig.tight_layout()
    fig.savefig(FIG / "fig4_confusion_matrices.png", dpi=180); fig.savefig(DOC_FIG / "fig4_confusion_matrices.png", dpi=180); plt.close(fig)

def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), fill); tcPr.append(shd)

def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr(); tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None: tcMar = OxmlElement("w:tcMar"); tcPr.append(tcMar)
    for m,v in (("top",top),("start",start),("bottom",bottom),("end",end)):
        node=tcMar.find(qn(f"w:{m}"))
        if node is None: node=OxmlElement(f"w:{m}"); tcMar.append(node)
        node.set(qn("w:w"),str(v)); node.set(qn("w:type"),"dxa")

def font_run(run, size=10.5, bold=False, color=None):
    run.font.name = "宋体"; rpr=run._element.get_or_add_rPr(); rpr.rFonts.set(qn("w:eastAsia"), "宋体")
    lang=rpr.find(qn("w:lang"))
    if lang is None: lang=OxmlElement("w:lang"); rpr.append(lang)
    lang.set(qn("w:val"),"zh-CN"); lang.set(qn("w:eastAsia"),"zh-CN")
    run.font.size = Pt(size); run.bold = bold
    if color: run.font.color.rgb = RGBColor(*color)

def format_para(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=0, line=1.5):
    p.alignment=align; p.paragraph_format.space_before=Pt(before); p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=line
    for r in p.runs: font_run(r)

def add_body(doc, text, bold_prefix=None):
    p=doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        font_run(p.add_run(bold_prefix),bold=True); font_run(p.add_run(text[len(bold_prefix):]))
    else: font_run(p.add_run(text))
    format_para(p); return p

def add_heading(doc, text, level=1):
    p=doc.add_paragraph(); font_run(p.add_run(text), size=14 if level==1 else 12, bold=True, color=(31,111,178))
    p.paragraph_format.space_before=Pt(8 if level==1 else 4); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.5
    p.paragraph_format.keep_with_next=True; return p

def add_table(doc, headers, rows, widths=None):
    table=doc.add_table(rows=1,cols=len(headers)); table.style="Table Grid"; table.autofit=False
    for i,h in enumerate(headers):
        c=table.rows[0].cells[i]; set_cell_shading(c,"DDEBF7"); c.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER; set_cell_margins(c)
        p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font_run(p.add_run(str(h)),9,bold=True); format_para(p,WD_ALIGN_PARAGRAPH.CENTER,line=1.0)
    for row in rows:
        cells=table.add_row().cells
        for i,val in enumerate(row):
            set_cell_margins(cells[i]); cells[i].vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=cells[i].paragraphs[0]; font_run(p.add_run(str(val)),9); format_para(p,WD_ALIGN_PARAGRAPH.CENTER,line=1.0)
    if widths:
        for row in table.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    return table

def add_figure(doc, path, caption, interpretation, width=15.8):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run(); r.add_picture(str(path),width=Cm(width)); p.paragraph_format.keep_with_next=True
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; font_run(c.add_run(caption),10.5,bold=True); format_para(c,WD_ALIGN_PARAGRAPH.CENTER)
    add_body(doc, "图表解读："+interpretation)

def build_docx():
    # Reuse Task 4's proven Word theme/font table, then replace its document body.
    doc=Document(ROOT/"TASK4"/"程冰晖 TASK4.docx")
    body=doc._element.body; sectPr=body.sectPr
    for child in list(body):
        if child is not sectPr: body.remove(child)
    sec=doc.sections[0]; sec.page_width=Cm(21); sec.page_height=Cm(29.7); sec.top_margin=Cm(2.54); sec.bottom_margin=Cm(2.54); sec.left_margin=Cm(2.54); sec.right_margin=Cm(2.54)
    styles=doc.styles
    normal=styles["Normal"]; normal.font.name="宋体"; nrpr=normal._element.get_or_add_rPr(); nrpr.rFonts.set(qn("w:eastAsia"),"宋体"); normal.font.size=Pt(10.5)
    nlang=nrpr.find(qn("w:lang"))
    if nlang is None: nlang=OxmlElement("w:lang"); nrpr.append(nlang)
    nlang.set(qn("w:val"),"zh-CN"); nlang.set(qn("w:eastAsia"),"zh-CN")
    normal.paragraph_format.line_spacing=1.5; normal.paragraph_format.space_before=Pt(0); normal.paragraph_format.space_after=Pt(0); normal.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    header=sec.header.paragraphs[0]; header.alignment=WD_ALIGN_PARAGRAPH.RIGHT; font_run(header.add_run("机器学习分类研究 · TASK5"),9,color=(102,115,133))
    footer=sec.footer.paragraphs[0]; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER; font_run(footer.add_run("程冰晖 · AI Quant Study"),9,color=(102,115,133))
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(24); p.paragraph_format.space_after=Pt(12); font_run(p.add_run("股票收益数据的分类机器学习研究"),18,bold=True,color=(23,32,51))
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; font_run(p.add_run("——逻辑回归、决策树、随机森林与 KNN 的训练及评价"),12,bold=True,color=(31,111,178)); format_para(p,WD_ALIGN_PARAGRAPH.CENTER)
    add_body(doc,"姓名：程冰晖    作业：TASK5    数据区间：2018-01-02 至 2026-07-10")
    add_body(doc,"摘要：本文使用三只已完成复权校验的股票日线数据构造下一交易日涨跌标签，以2018—2022年为训练集、2023—2024年为验证集、2025年至最近交易日为测试集，分别训练逻辑回归、决策树、随机森林和KNN分类模型，并通过混淆矩阵、ROC曲线与AUC评价样本外分类能力。结果表明，不同股票和模型的表现存在明显差异，AUC应结合混淆矩阵、精确率、召回率和F1综合解释。")

    add_heading(doc,"一、分类型机器学习算法",1)
    add_heading(doc,"（一）逻辑回归",2); add_body(doc,"逻辑回归通过线性组合计算样本属于上涨类别的概率，再使用Sigmoid函数将结果映射到0—1。模型系数可解释特征对上涨概率的方向与相对影响，优点是结构简单、训练稳定、便于作为基线；局限是决策边界主要为线性，难以自动捕捉复杂非线性关系。")
    add_heading(doc,"（二）决策树",2); add_body(doc,"决策树按照特征阈值反复划分样本，使子节点中的类别更纯。它能够表达非线性关系和特征交互，且规则直观；但深度过大时容易记住训练噪声，因此本实验将max_depth作为可配置参数并用验证集控制复杂度。")
    add_heading(doc,"（三）随机森林",2); add_body(doc,"随机森林在不同自助样本和随机特征子集上训练多棵决策树，再通过投票得到最终类别及概率。集成平均通常比单棵树更稳定，能够降低方差。本实验默认n_estimators=100、max_depth=10；树数、深度和类别权重均可调整。")
    add_heading(doc,"（四）KNN",2); add_body(doc,"K近邻根据标准化特征空间中距离最近的K个训练样本投票。它无需显式假设函数形式，但对特征尺度、K值和高维噪声较敏感，因此实验使用StandardScaler，并将n_neighbors设为可配置参数。")
    add_table(doc,["算法家族","分类实现","主要参数","主要特点"],[["逻辑回归","LogisticRegression","max_iter、C","可解释线性基线"],["决策树","DecisionTreeClassifier","max_depth","非线性、规则直观"],["随机森林","RandomForestClassifier","n_estimators、max_depth","集成稳定、可给重要性"],["KNN","KNeighborsClassifier","n_neighbors、weights","基于距离、需标准化"]],[3,4,4,5])

    add_heading(doc,"二、模型评价指标",1)
    add_heading(doc,"（一）混淆矩阵",2); add_body(doc,"混淆矩阵由真正例TP、假正例FP、真负例TN和假负例FN组成。股票方向预测中，TP表示真实上涨且预测上涨，FP表示真实下跌但预测上涨。由此可计算准确率、精确率、召回率和F1。")
    add_heading(doc,"（二）ROC曲线",2); add_body(doc,"ROC曲线在所有分类阈值下绘制真正率TPR与假正率FPR的关系。曲线越靠近左上角，模型在不同阈值下区分上涨与下跌的能力通常越强；对角线对应接近随机排序。")
    add_heading(doc,"（三）AUC",2); add_body(doc,"AUC是ROC曲线下面积，可理解为随机抽取一个正类和一个负类时，模型将正类概率排得更高的可能性。AUC=0.5通常表示接近随机，越接近1区分能力越强。AUC不等同于交易收益，也不能替代对类别分布和交易成本的分析。")

    add_heading(doc,"三、Python实现方案",1)
    add_body(doc,"1. 加载数据：读取复权日线，按股票代码和交易日期排序，检查OHLC关系、重复记录与缺失值。")
    add_body(doc,"2. 构造标签：以未来一个交易日复权收盘收益率是否大于0构造target_up。")
    add_body(doc,"3. 构造特征：使用1/5/10/20日收益、均线偏离、波动率、ATR、RSI、MACD、振幅和成交量相对均量等22个特征。")
    add_body(doc,"4. 数据划分：2018—2022年训练、2023—2024年验证、2025年至今测试；预处理只在训练数据拟合，不随机打乱。")
    add_body(doc,"5. 模型训练与测试：在训练/验证阶段确定方案后，用训练集与验证集合并重训，最后一次性评估测试集。")
    add_figure(doc,DOC_FIG/"fig1_time_split.png","图1 三只股票训练集、验证集与测试集时间划分","所有划分都沿时间顺序进行。中芯国际A股因2020年上市，训练样本从实际可用日期开始；测试集始终保持隔离。")

    add_heading(doc,"四、实验结果与分析",1)
    test_auc=METRICS[(METRICS.dataset=="test")&(METRICS.metric=="AUC")]
    rows=[]
    for symbol in LABELS:
        for model in ["logistic_regression","decision_tree_classifier","random_forest_classifier","knn_classifier"]:
            v=test_auc[(test_auc.symbol==symbol)&(test_auc.model==model)].value.iloc[0]
            rows.append([LABELS[symbol],MODEL_CN[model],num(v),int(test_auc[(test_auc.symbol==symbol)&(test_auc.model==model)].n_samples.iloc[0])])
    add_table(doc,["股票","模型","测试集AUC","测试样本数"],rows,[4,4,3,3])
    add_figure(doc,DOC_FIG/"fig2_auc_comparison.png","图2 测试集分类模型AUC对比","AUC结果显示模型效果随股票而变化，整体优势有限，说明短周期涨跌具有较强噪声。应避免只选择单一股票或单一指标得出过度结论。")
    add_figure(doc,DOC_FIG/"fig3_roc_random_forest.png","图3 随机森林测试集ROC曲线","三条ROC曲线以0.5对角线为参照。曲线形状展示了阈值从高到低变化时召回更多上涨日与增加误报之间的权衡。")
    add_figure(doc,DOC_FIG/"fig4_confusion_matrices.png","图4 随机森林测试集混淆矩阵","混淆矩阵揭示模型是否偏向预测某一类别。即便AUC相近，固定0.5阈值下的TP、FP、TN和FN结构也可能明显不同。")

    add_heading(doc,"五、结论与局限",1)
    add_body(doc,"本文完成了股票收益方向分类的完整流程，包括数据加载、标签与特征构造、时间序列划分、四种分类模型训练，以及混淆矩阵、ROC和AUC评价。逻辑回归提供可解释基线，决策树表达非线性规则，随机森林通过集成降低方差，KNN则提供基于邻近样本的比较。")
    add_body(doc,"需要强调的是，分类指标不等同于可交易收益。本实验尚未把手续费、滑点和持仓规则纳入标签与模型评价；日线样本还存在制度变化、公司行动和市场状态漂移。后续可加入滚动训练、概率阈值优化、类别权重、交易成本标签及更多股票，检验模型稳定性。")
    add_heading(doc,"附录：关键参数",1)
    add_table(doc,["类别","参数","默认值","说明"],[["标签","forecast_horizon","1","预测下一交易日"],["分类","threshold","0.0 / 概率0.5","收益标签阈值/概率决策阈值"],["随机森林","n_estimators","100","树数量"],["随机森林/决策树","max_depth","10","最大深度"],["KNN","n_neighbors","5","邻居数量"],["时间划分","train/validation/test","2018—22/2023—24/2025—","固定年份留出"]],[3,4,4,6])
    doc.core_properties.title="股票收益数据的分类机器学习研究"; doc.core_properties.author="程冰晖"; doc.save(DOCX)

def build_html():
    cls=PRED[(PRED.task_type=="classification")].copy(); cls["date"]=cls.date.dt.strftime("%Y-%m-%d")
    metric_records=METRICS[METRICS.task_type.eq("classification")].replace({np.nan:None}).to_dict("records")
    payload={"predictions":cls.replace({np.nan:None}).to_dict("records"),"metrics":metric_records,"symbols":[{"code":s,"label":LABELS[s]} for s in LABELS],"models":[{"key":m,"label":MODEL_CN[m]} for m in ["logistic_regression","decision_tree_classifier","random_forest_classifier","knn_classifier"]]}
    template=(TASK/"task5_report_template.html").read_text(encoding="utf-8")
    HTML.write_text(template.replace("__PAYLOAD__",json.dumps(payload,ensure_ascii=False,separators=(",",":"))),encoding="utf-8")

if __name__=="__main__":
    make_figures(); build_docx(); build_html(); print(DOCX); print(HTML)
